"""æ–‡ä»¶é¢„è§ˆæœåŠ¡ - æ”¯æŒ Excelã€Wordã€PPTã€PDF ç­‰æ–‡æ¡£çš„åŸç”Ÿé¢„è§ˆ"""
import os
import io
import html
import hashlib
import tempfile
from pathlib import Path
from typing import Optional, Tuple
import base64


# é¢„è§ˆç¼“å­˜ç›®å½•
_preview_cache_dir: Optional[Path] = None


def get_preview_cache_dir() -> Path:
    """è·å–é¢„è§ˆç¼“å­˜ç›®å½•"""
    global _preview_cache_dir
    if _preview_cache_dir is None:
        _preview_cache_dir = Path(tempfile.gettempdir()) / "webrpa_file_preview"
        _preview_cache_dir.mkdir(exist_ok=True)
    return _preview_cache_dir


def get_cache_key(file_path: Path) -> str:
    """ç”Ÿæˆæ–‡ä»¶ç¼“å­˜é”®"""
    file_hash = hashlib.md5(str(file_path).encode()).hexdigest()
    mtime = int(file_path.stat().st_mtime)
    return f"{file_hash}_{mtime}"


def preview_excel(file_path: Path) -> Tuple[str, str]:
    """
    é¢„è§ˆ Excel æ–‡ä»¶ï¼Œè¿”å› (html_content, content_type)
    æ”¯æŒ .xlsx, .xls, .csv
    """
    ext = file_path.suffix.lower()
    
    try:
        if ext == '.csv':
            return _preview_csv(file_path)
        elif ext == '.xlsx':
            return _preview_xlsx(file_path)
        elif ext == '.xls':
            return _preview_xls(file_path)
        else:
            return f"<p>ä¸æ”¯æŒçš„ Excel æ ¼å¼: {ext}</p>", "text/html"
    except Exception as e:
        return f"<p>é¢„è§ˆå¤±è´¥: {html.escape(str(e))}</p>", "text/html"


def _preview_csv(file_path: Path) -> Tuple[str, str]:
    """é¢„è§ˆ CSV æ–‡ä»¶"""
    import csv
    
    rows = []
    encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16', 'latin-1']
    
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding, newline='') as f:
                reader = csv.reader(f)
                rows = list(reader)[:1000]  # é™åˆ¶è¡Œæ•°
            break
        except (UnicodeDecodeError, UnicodeError):
            continue
    
    if not rows:
        return "<p>æ— æ³•è¯»å– CSV æ–‡ä»¶</p>", "text/html"
    
    return _generate_table_html(rows, file_path.name), "text/html"


def _preview_xlsx(file_path: Path) -> Tuple[str, str]:
    """é¢„è§ˆ XLSX æ–‡ä»¶"""
    import openpyxl
    
    wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
    sheets_html = []
    
    for sheet_name in wb.sheetnames[:10]:  # é™åˆ¶å·¥ä½œè¡¨æ•°é‡
        ws = wb[sheet_name]
        rows = []
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i >= 1000:  # é™åˆ¶è¡Œæ•°
                break
            rows.append(list(row))
        
        if rows:
            sheet_html = f'<div class="sheet"><h3>{html.escape(sheet_name)}</h3>'
            sheet_html += _generate_table_html(rows, None, show_title=False)
            sheet_html += '</div>'
            sheets_html.append(sheet_html)
    
    wb.close()
    
    content = _get_preview_wrapper(
        file_path.name,
        '<div class="sheets">' + ''.join(sheets_html) + '</div>',
        extra_style='.sheet { margin-bottom: 24px; } .sheet h3 { margin-bottom: 12px; color: #3b82f6; }'
    )
    return content, "text/html"


def _preview_xls(file_path: Path) -> Tuple[str, str]:
    """é¢„è§ˆ XLS æ–‡ä»¶"""
    import xlrd
    
    wb = xlrd.open_workbook(file_path)
    sheets_html = []
    
    for sheet_idx in range(min(wb.nsheets, 10)):
        ws = wb.sheet_by_index(sheet_idx)
        rows = []
        for row_idx in range(min(ws.nrows, 1000)):
            row = [ws.cell_value(row_idx, col_idx) for col_idx in range(ws.ncols)]
            rows.append(row)
        
        if rows:
            sheet_html = f'<div class="sheet"><h3>{html.escape(ws.name)}</h3>'
            sheet_html += _generate_table_html(rows, None, show_title=False)
            sheet_html += '</div>'
            sheets_html.append(sheet_html)
    
    content = _get_preview_wrapper(
        file_path.name,
        '<div class="sheets">' + ''.join(sheets_html) + '</div>',
        extra_style='.sheet { margin-bottom: 24px; } .sheet h3 { margin-bottom: 12px; color: #3b82f6; }'
    )
    return content, "text/html"


def _generate_table_html(rows: list, filename: Optional[str], show_title: bool = True) -> str:
    """ç”Ÿæˆè¡¨æ ¼ HTML"""
    if not rows:
        return "<p>ç©ºè¡¨æ ¼</p>"
    
    table_html = '<div class="table-wrapper"><table>'
    
    # è¡¨å¤´
    if rows:
        table_html += '<thead><tr>'
        for cell in rows[0]:
            cell_str = str(cell) if cell is not None else ''
            table_html += f'<th>{html.escape(cell_str)}</th>'
        table_html += '</tr></thead>'
    
    # è¡¨ä½“
    table_html += '<tbody>'
    for row in rows[1:]:
        table_html += '<tr>'
        for cell in row:
            cell_str = str(cell) if cell is not None else ''
            table_html += f'<td>{html.escape(cell_str)}</td>'
        table_html += '</tr>'
    table_html += '</tbody></table></div>'
    
    if show_title and filename:
        return _get_preview_wrapper(filename, table_html)
    return table_html


def preview_word(file_path: Path) -> Tuple[str, str]:
    """
    é¢„è§ˆ Word æ–‡ä»¶ï¼Œè¿”å› (html_content, content_type)
    æ”¯æŒ .docx
    """
    ext = file_path.suffix.lower()
    
    try:
        if ext == '.docx':
            return _preview_docx(file_path)
        elif ext == '.doc':
            return "<p>æš‚ä¸æ”¯æŒ .doc æ ¼å¼ï¼Œè¯·è½¬æ¢ä¸º .docx</p>", "text/html"
        else:
            return f"<p>ä¸æ”¯æŒçš„ Word æ ¼å¼: {ext}</p>", "text/html"
    except Exception as e:
        return f"<p>é¢„è§ˆå¤±è´¥: {html.escape(str(e))}</p>", "text/html"


def _preview_docx(file_path: Path) -> Tuple[str, str]:
    """é¢„è§ˆ DOCX æ–‡ä»¶"""
    from docx import Document
    from docx.shared import Inches
    
    doc = Document(file_path)
    content_parts = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            # æ£€æŸ¥æ®µè½æ ·å¼
            style_name = para.style.name if para.style else ''
            if 'Heading' in style_name or 'heading' in style_name.lower() or 'æ ‡é¢˜' in style_name:
                level = 2
                if '1' in style_name:
                    level = 1
                elif '2' in style_name:
                    level = 2
                elif '3' in style_name:
                    level = 3
                content_parts.append(f'<h{level}>{html.escape(para.text)}</h{level}>')
            elif 'Title' in style_name or 'title' in style_name.lower():
                content_parts.append(f'<h1 class="doc-title">{html.escape(para.text)}</h1>')
            else:
                content_parts.append(f'<p>{html.escape(para.text)}</p>')
    
    # å¤„ç†è¡¨æ ¼
    for table in doc.tables:
        table_html = '<table class="doc-table">'
        for i, row in enumerate(table.rows):
            table_html += '<tr>'
            tag = 'th' if i == 0 else 'td'
            for cell in row.cells:
                table_html += f'<{tag}>{html.escape(cell.text)}</{tag}>'
            table_html += '</tr>'
        table_html += '</table>'
        content_parts.append(table_html)
    
    if not content_parts:
        content_parts.append('<p class="empty-doc">æ–‡æ¡£å†…å®¹ä¸ºç©º</p>')
    
    content = _get_preview_wrapper(
        file_path.name,
        '<div class="doc-content">' + ''.join(content_parts) + '</div>',
        extra_style='''
            .doc-content { max-width: 800px; margin: 0 auto; line-height: 1.8; }
            .doc-content p { margin-bottom: 12px; text-indent: 2em; }
            .doc-content h1, .doc-content h2, .doc-content h3 { margin: 20px 0 12px; text-indent: 0; }
            .doc-content .doc-title { text-align: center; font-size: 24px; margin-bottom: 24px; }
            .doc-table { width: 100%; margin: 16px 0; border-collapse: collapse; }
            .doc-table th, .doc-table td { padding: 10px 12px; border: 1px solid #3f3f46; text-indent: 0; }
            .doc-table th { background: #27272a; font-weight: 600; }
            .empty-doc { text-align: center; color: #71717a; }
        '''
    )
    return content, "text/html"


def preview_ppt(file_path: Path) -> Tuple[str, str]:
    """
    é¢„è§ˆ PPT æ–‡ä»¶ï¼Œè¿”å› (html_content, content_type)
    æ”¯æŒ .pptx
    """
    ext = file_path.suffix.lower()
    
    try:
        if ext == '.pptx':
            return _preview_pptx(file_path)
        elif ext == '.ppt':
            return "<p>æš‚ä¸æ”¯æŒ .ppt æ ¼å¼ï¼Œè¯·è½¬æ¢ä¸º .pptx</p>", "text/html"
        else:
            return f"<p>ä¸æ”¯æŒçš„ PPT æ ¼å¼: {ext}</p>", "text/html"
    except ImportError:
        return "<p>é¢„è§ˆ PPT éœ€è¦å®‰è£… python-pptx åº“: pip install python-pptx</p>", "text/html"
    except Exception as e:
        return f"<p>é¢„è§ˆå¤±è´¥: {html.escape(str(e))}</p>", "text/html"


def _preview_pptx(file_path: Path) -> Tuple[str, str]:
    """é¢„è§ˆ PPTX æ–‡ä»¶ - ä½¿ç”¨ zipfile ç›´æ¥è§£æ XML"""
    import zipfile
    import xml.etree.ElementTree as ET
    
    slides_html = []
    
    # PPTX æ˜¯ä¸€ä¸ª ZIP æ–‡ä»¶ï¼Œç›´æ¥è§£æ XML
    with zipfile.ZipFile(file_path, 'r') as zf:
        # è·å–æ‰€æœ‰å¹»ç¯ç‰‡æ–‡ä»¶
        slide_files = sorted([f for f in zf.namelist() if f.startswith('ppt/slides/slide') and f.endswith('.xml')])
        
        for idx, slide_file in enumerate(slide_files[:50]):
            slide_content = []
            slide_content.append(f'<div class="slide-number">ç¬¬ {idx + 1} é¡µ</div>')
            
            try:
                with zf.open(slide_file) as f:
                    tree = ET.parse(f)
                    root = tree.getroot()
                    
                    # æå–æ‰€æœ‰æ–‡æœ¬
                    texts = []
                    for elem in root.iter():
                        # æŸ¥æ‰¾æ–‡æœ¬å…ƒç´  (a:t)
                        if elem.tag.endswith('}t') and elem.text:
                            texts.append(elem.text)
                    
                    # åˆå¹¶è¿ç»­æ–‡æœ¬ï¼ŒæŒ‰æ®µè½åˆ†ç»„
                    if texts:
                        # ç¬¬ä¸€ä¸ªé€šå¸¸æ˜¯æ ‡é¢˜
                        if len(texts) > 0:
                            slide_content.append(f'<h2>{html.escape(texts[0])}</h2>')
                        # å…¶ä½™æ˜¯å†…å®¹
                        for text in texts[1:]:
                            if text.strip():
                                slide_content.append(f'<p>{html.escape(text)}</p>')
            except Exception:
                pass
            
            if len(slide_content) == 1:
                slide_content.append('<p class="empty-slide">ï¼ˆæ­¤é¡µæ— æ–‡æœ¬å†…å®¹ï¼‰</p>')
            
            slides_html.append(f'<div class="slide">{"".join(slide_content)}</div>')
    
    content = _get_preview_wrapper(
        file_path.name,
        '<div class="slides">' + ''.join(slides_html) + '</div>',
        extra_style='''
            .slides { display: flex; flex-direction: column; gap: 24px; }
            .slide { background: #27272a; border-radius: 8px; padding: 24px; border: 1px solid #3f3f46; }
            .slide-number { font-size: 12px; color: #71717a; margin-bottom: 12px; }
            .slide h2 { font-size: 20px; margin-bottom: 16px; color: #3b82f6; }
            .slide p { margin-bottom: 8px; line-height: 1.6; }
            .empty-slide { color: #71717a; font-style: italic; }
        '''
    )
    return content, "text/html"


def preview_pdf_as_images(file_path: Path) -> Tuple[str, str]:
    """
    å°† PDF è½¬æ¢ä¸ºå›¾ç‰‡é¢„è§ˆï¼Œè¿”å› (html_content, content_type)
    ä½¿ç”¨ pdf2image åº“æ›¿ä»£ PyMuPDF
    """
    try:
        from pdf2image import convert_from_path
        
        # è·å– backend ç›®å½•çš„è·¯å¾„
        backend_root = Path(__file__).parent.parent.parent
        poppler_path = backend_root / 'poppler' / 'Library' / 'bin'
        
        # è½¬æ¢ PDF ä¸ºå›¾ç‰‡ï¼Œé™åˆ¶é¡µæ•°
        if poppler_path.exists():
            images = convert_from_path(str(file_path), dpi=150, first_page=1, last_page=50, poppler_path=str(poppler_path))
        else:
            # å¦‚æœå†…ç½® poppler ä¸å­˜åœ¨ï¼Œå°è¯•ä½¿ç”¨ç³»ç»Ÿ PATH ä¸­çš„ poppler
            images = convert_from_path(str(file_path), dpi=150, first_page=1, last_page=50)
        
        pages_html = []
        
        for page_num, img in enumerate(images):
            # å°†å›¾ç‰‡è½¬æ¢ä¸º base64
            img_buffer = io.BytesIO()
            img.save(img_buffer, format='PNG')
            img_bytes = img_buffer.getvalue()
            b64_image = base64.b64encode(img_bytes).decode('utf-8')
            
            pages_html.append(f'''
                <div class="pdf-page">
                    <div class="page-number">ç¬¬ {page_num + 1} é¡µ</div>
                    <img src="data:image/png;base64,{b64_image}">
                </div>
            ''')
        
        content = _get_preview_wrapper(
            file_path.name,
            '<div class="pdf-pages">' + ''.join(pages_html) + '</div>',
            extra_style='''
                .pdf-pages { display: flex; flex-direction: column; align-items: center; gap: 16px; }
                .pdf-page { background: #27272a; border-radius: 8px; padding: 16px; text-align: center; }
                .page-number { font-size: 12px; color: #71717a; margin-bottom: 8px; }
                .pdf-page img { max-width: 100%; border-radius: 4px; box-shadow: 0 4px 12px rgba(0,0,0,0.3); }
            '''
        )
        return content, "text/html"
    except ImportError:
        error_html = "<p>PDF é¢„è§ˆéœ€è¦å®‰è£… pdf2image åº“: pip install pdf2image</p>"
        return error_html, "text/html"
    except Exception as e:
        error_msg = str(e)
        if "poppler" in error_msg.lower():
            error_html = "<div style='padding: 20px; color: #fafafa;'>"
            error_html += "<h2 style='color: #ef4444;'>PDF é¢„è§ˆéœ€è¦ poppler å·¥å…·</h2>"
            error_html += "<p>è¯·å°† poppler æ–‡ä»¶å¤¹æ”¾ç½®åˆ° backend ç›®å½•ä¸‹</p>"
            error_html += "<p>ä¸‹è½½åœ°å€: <a href='https://github.com/oschwartz10612/poppler-windows/releases' target='_blank' style='color: #3b82f6;'>https://github.com/oschwartz10612/poppler-windows/releases</a></p>"
            error_html += f"<p style='color: #71717a; margin-top: 20px;'>é”™è¯¯è¯¦æƒ…: {html.escape(error_msg)}</p>"
            error_html += "</div>"
            return error_html, "text/html"
        return f"<p>PDF é¢„è§ˆå¤±è´¥: {html.escape(error_msg)}</p>", "text/html"


def _get_preview_wrapper(filename: str, content: str, extra_style: str = '') -> str:
    """ç”Ÿæˆé¢„è§ˆé¡µé¢åŒ…è£…å™¨"""
    return f'''<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{html.escape(filename)} - é¢„è§ˆ</title>
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{ 
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; 
            background: #18181b; 
            color: #fafafa; 
            padding: 20px;
            min-height: 100vh;
        }}
        .header {{
            text-align: center;
            margin-bottom: 24px;
            padding-bottom: 16px;
            border-bottom: 1px solid #27272a;
        }}
        .header h1 {{
            font-size: 18px;
            font-weight: 600;
            color: #fafafa;
            word-break: break-all;
        }}
        .table-wrapper {{
            overflow-x: auto;
            margin: 16px 0;
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
            font-size: 13px;
        }}
        th, td {{
            padding: 10px 12px;
            text-align: left;
            border: 1px solid #3f3f46;
            white-space: nowrap;
            max-width: 300px;
            overflow: hidden;
            text-overflow: ellipsis;
        }}
        th {{
            background: #27272a;
            font-weight: 600;
            position: sticky;
            top: 0;
        }}
        tr:hover td {{
            background: #27272a;
        }}
        {extra_style}
    </style>
</head>
<body>
    <div class="header">
        <h1>ğŸ“„ {html.escape(filename)}</h1>
    </div>
    {content}
</body>
</html>'''


def get_preview_content(file_path: Path) -> Optional[Tuple[bytes, str]]:
    """
    è·å–æ–‡ä»¶é¢„è§ˆå†…å®¹
    è¿”å› (content_bytes, content_type) æˆ– None
    """
    ext = file_path.suffix.lower()
    
    # Excel æ–‡ä»¶
    if ext in ['.xlsx', '.xls', '.csv']:
        content, content_type = preview_excel(file_path)
        return content.encode('utf-8'), content_type
    
    # Word æ–‡ä»¶
    if ext in ['.docx', '.doc']:
        content, content_type = preview_word(file_path)
        return content.encode('utf-8'), content_type
    
    # PPT æ–‡ä»¶
    if ext in ['.pptx', '.ppt']:
        content, content_type = preview_ppt(file_path)
        return content.encode('utf-8'), content_type
    
    return None


def is_previewable_document(filename: str) -> bool:
    """æ£€æŸ¥æ–‡ä»¶æ˜¯å¦æ”¯æŒæ–‡æ¡£é¢„è§ˆ"""
    ext = Path(filename).suffix.lower()
    return ext in ['.xlsx', '.xls', '.csv', '.docx', '.doc', '.pptx', '.ppt']
